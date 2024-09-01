# More input prompt questions
# Any short explaination about the dataset?
# Any prediction involved

import pandas as pd
import numpy as np
import os
from statsmodels.tools import add_constant
from statsmodels.stats.outliers_influence import variance_inflation_factor
from openai import OpenAI
from jinja2 import Environment, FileSystemLoader
from markupsafe import escape, Markup
from docx import Document

def nl2br(value):
    # Split the text into lines and wrap each line with a list item tag
    lines = value.split('\n')
    formatted_lines = ''.join([f'<li>{line.strip()}</li>' for line in lines if line.strip() != ''])
    return Markup(f'<ul>{formatted_lines}</ul>')

def load_dataset(csv_file):
    return pd.read_csv(csv_file, dtype={'col_name': 'float32'})  # Example downcasting

def mad_based_outlier_detection(data, threshold=3.5):
    """
    Detect outliers using the Median Absolute Deviation (MAD) method.
    :param data: Pandas Series of numeric values.
    :param threshold: Threshold for defining outliers.
    :return: Boolean Series indicating outliers.
    """
    median = np.median(data)
    mad = np.median(np.abs(data - median))
    modified_z_score = 0.6745 * (data - median) / mad if mad else 0
    return np.abs(modified_z_score) > threshold

def perform_eda(df):
    
    shape = df.shape
    
    # Numeric Columns EDA
    numeric_df = df.select_dtypes(include=['number'])

    # Correlation matrix filtering
    corr_matrix = numeric_df.corr()
    filtered_corr = corr_matrix[(corr_matrix > 0.3) | (corr_matrix < -0.3)]
    filtered_corr = filtered_corr.where(np.abs(filtered_corr) != 1)

    # VIF Calculation
    # Fill NaN values with the median for each column and add a constant for the intercept term
    X = add_constant(numeric_df.apply(lambda x: x.fillna(x.median()), axis=0))

    # Calculate VIF for each variable
    vif_df = pd.DataFrame()
    vif_df["Variable"] = X.columns
    vif_df["VIF"] = [round(variance_inflation_factor(X.values, i), 2) for i in range(X.shape[1])]

    # Exclude the constant term from the VIF dictionary
    vif_df = vif_df[vif_df['Variable'] != 'const']

    # Create a dictionary with variables as keys and rounded VIFs as values
    vif_dict = dict(zip(vif_df['Variable'], vif_df['VIF']))

    numeric_cols = {}
    unique_threshold = 10
    for col in numeric_df.columns:
        stats = {}
        
        # Only check columns with a large variety of values
        unique_values_count = numeric_df[col].nunique()
        if unique_values_count > unique_threshold:
            # Detect outliers using MAD
            outliers = mad_based_outlier_detection(numeric_df[col].dropna())
            if outliers.any():
                stats['Outlier'] = "Yes"
        
        # Include Null Count if there are missing values
        null_count = numeric_df[col].isnull().sum()
        if null_count > 0:
            stats['null_count'] = null_count
        
        # Include Mean and Median by default
        stats['mean'] = numeric_df[col].mean().round(2)
        stats['median'] = numeric_df[col].median().round(2)
        
        # Include Skew if significant
        skew = numeric_df[col].skew().round(2)
        if abs(skew) > 0.5:
            stats['skew'] = skew
        
        # Include Kurtosis if it's unusually high or low
        kurtosis = numeric_df[col].kurt().round(2)
        if abs(kurtosis) > 3:
            stats['kurtosis'] = kurtosis
        
        # Include VIF if it indicates multicollinearity
        vif_value = vif_dict.get(col)
        if vif_value and vif_value > 5:
            stats['vif'] = vif_value
        
        # Include correlations only if they are strong
        correlations = filtered_corr[col].round(2).dropna().to_dict()
        if correlations:
            stats['corr'] = correlations
        
        # Include Top Value Counts if a single value dominates
        value_counts = numeric_df[col].value_counts(normalize=True).head(3) * 100
        value_counts = value_counts.round(2).astype(str) + "%"
        
        if any(float(str(percent).strip('%')) > 50 for percent in value_counts):
            stats['value_count'] = value_counts.to_dict()
        
        if stats:
            numeric_cols[col] = stats

    # String Columns EDA
    string_df = df.select_dtypes(include='object')
    string_cols = {}
    for col in string_df.columns:
        stats = {}
        
        # Add null_count only if there are null values
        null_count = string_df[col].isnull().sum()
        if null_count > 0:
            stats['null_count'] = null_count
        
        # Add unique count
        stats['unique_count'] = string_df[col].dropna().nunique()
        
        # Add average token length
        stats['avg_len'] = int(string_df[col].dropna().str.split().str.len().mean())
        
        # Add top 3 value counts if they are significant
        if string_df[col].dropna().nunique() < 0.1 * len(string_df):
            top_3_value_counts = string_df[col].dropna().value_counts(normalize=True).head(3) * 100  # Convert to percentages
            top_3_value_counts = top_3_value_counts.round(2).astype(str) + "%"  # Convert to string and append '%'
            stats['top_3_value_counts'] = top_3_value_counts.to_dict()
        else:
            stats['top_3_value_counts'] = "no notable value counts"
        
        # Add 3 random values, truncated to the first 20 tokens
        stats['3_random_values'] = [
            value[:20] for value in string_df[col].dropna().sample(3).to_list()
        ]
        
        # Add the stats to the string_cols dictionary
        string_cols[col] = stats

    return {'numeric_cols': numeric_cols, 'string_cols': string_cols, 'shape':shape}

def interact_with_chatgpt(eda_results, target_column, api_key):
    client = OpenAI(api_key=api_key)

    # Define the first group of tasks (1-5) with formatting instructions and specificity request
    prompt_group1 = f"""
    I performed EDA on a dataset with the following details:

    Shape: {eda_results['shape']}
    Numeric Columns: {eda_results['numeric_cols']}
    String Columns: {eda_results['string_cols']}
    Target Column: {target_column}

    Please provide detailed and specific insights for the following, focusing on potential problems and actionable steps derived from the EDA information:
    1. Critical data cleaning steps.
    2. Issues with data distribution.
    3. Feature engineering and data processing suggestions.
    4. Grouping or aggregation ideas to uncover insights.
    5. Most valuable predictive columns and reasoning.

    Format the response in HTML using:
    - <h3> for main points,
    - <ul> and <li> for lists,
    - <p> for explanations,
    - Proper HTML nesting.

    Ensure the HTML is ready for direct insertion into a document.
    """

    # Define the second group of tasks (6-9) with formatting instructions and specificity request
    prompt_group2 = f"""
    I performed EDA on a dataset with the following details:

    Shape: {eda_results['shape']}
    Numeric Columns: {eda_results['numeric_cols']}
    String Columns: {eda_results['string_cols']}
    Target Column: {target_column}

    Please provide detailed and specific insights for the following, focusing on potential problems and actionable steps derived from the EDA information:
    6. Appropriate ML models and justification.
    7. Model assumptions that need to be checked and data to be transformed for the most fit model.
    8. Suggested success metrics for the project and models.
    9. A summary roadmap for the project.

    Format the response in HTML using:
    - <h3> for main points,
    - <ul> and <li> for lists,
    - <p> for explanations,
    - Proper HTML nesting.

    Ensure the HTML is ready for direct insertion into a document.
    """

    # Make API call for the first group of tasks
    response_group1 = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt_group1}],
        max_tokens=1500  # Allocate up to 1500 tokens for the output
    )

    # Make API call for the second group of tasks
    response_group2 = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt_group2}],
        max_tokens=1500  # Allocate up to 1500 tokens for the output
    )

    # Combine the responses
    combined_response = f"{response_group1.choices[0].message.content}\n\n{response_group2.choices[0].message.content}"

    # Perform a final enhancement call to address any remaining high-level content or lack of specificity
    prompt_enhancement = f"""
    The combined response provided was still a bit high-level and not entirely specific to the dataset. 
    Please review the combined response and give more detailed and specific recommendations, 
    especially focusing on potential problems and actionable steps that can be derived from the EDA information provided.
    Here is the combined response:

    {combined_response}

    Ensure the final response is formatted in HTML using:
    - <h3> for main points,
    - <ul> and <li> for lists,
    - <p> for explanations,
    - Proper HTML nesting.

    Ensure the HTML is ready for direct insertion into a document.
    """

    enhanced_response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt_enhancement}],
        max_tokens=3000
    )

    # The enhanced response should already be formatted HTML, so we can return it directly
    return enhanced_response.choices[0].message.content



def create_report_directory():
    base_dir = "EDA_Report"
    counter = 1

    while True:
        report_dir = f"{base_dir}_{counter}"
        if not os.path.exists(report_dir):
            os.makedirs(report_dir)
            return report_dir
        counter += 1

def generate_html_report(eda_results, suggestions, output_html_file):
    # Create Jinja2 environment and register the nl2br filter
    env = Environment(loader=FileSystemLoader('.'))
    env.filters['nl2br'] = nl2br  # Register the custom filter

    # Load the template
    template = env.get_template('report_template.html')

    # Render the template with provided data
    html_content = template.render(eda_results=eda_results, suggestions=suggestions)

    # Write the generated HTML to a file
    with open(output_html_file, 'w') as f:
        f.write(html_content)

def generate_docx_report(eda_results, suggestions, output_docx_file):
    doc = Document()
    doc.add_heading('Exploratory Data Analysis Report', 0)

    doc.add_heading('Numeric Columns Analysis', level=1)
    for col, details in eda_results['numeric_cols'].items():
        doc.add_heading(col, level=2)
        for key, value in details.items():
            doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('String Columns Analysis', level=1)
    for col, details in eda_results['string_cols'].items():
        doc.add_heading(col, level=2)
        for key, value in details.items():
            doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('ChatGPT Suggestions', level=1)
    doc.add_paragraph(suggestions)

    doc.save(output_docx_file)

if __name__ == "__main__":
    csv_file = "train_1.csv"
    target_column = "Purcahse"
    my_api_key = "..."
    additional_information = input("Any additional information you want ChatGPT to know?")

    # Create a new report directory
    report_dir = create_report_directory()

    # Define file paths
    output_html_file = os.path.join(report_dir, f'{report_dir}.html')
    output_docx_file = os.path.join(report_dir, f'{report_dir}.docx')

    df = load_dataset(csv_file)
    eda_results = perform_eda(df)
    suggestions = interact_with_chatgpt(eda_results, target_column, my_api_key)

    # Generate reports
    generate_html_report(eda_results, suggestions, output_html_file)
    generate_docx_report(eda_results, suggestions, output_docx_file)

    print(f"EDA and report generation completed. Reports saved in {report_dir}.")
